"""Convert results/MANUSCRIPT.md into a reviewable Word document with embedded
figures and tables. Tailored to this manuscript's markdown structure."""
import os, re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image

SRC = "results/MANUSCRIPT.md"
OUT = "results/MANUSCRIPT.docx"
FIGDIR = "results"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(1))
nf = doc.styles["Normal"].font
nf.name, nf.size = "Calibri", Pt(10.5)
MAXW = 6.4


def add_runs(par, text):
    text = text.replace("\\*", "*")
    for tok in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not tok:
            continue
        if tok.startswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*"):
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)


def add_image(path):
    full = os.path.join(FIGDIR, path)
    if not os.path.exists(full):
        doc.add_paragraph("[missing figure: %s]" % path)
        return
    w, h = Image.open(full).size
    width = MAXW
    if width * h / w > 8.4:
        width = 8.4 * w / h
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(full, width=Inches(width))


def add_table(block):
    rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in block]
    rows = [r for r in rows if not all(re.match(r"^:?-+:?$", c) for c in r)]
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            p = t.cell(ri, ci).paragraphs[0]
            add_runs(p, row[ci] if ci < len(row) else "")
            for run in p.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
    doc.add_paragraph()


lines = open(SRC, encoding="utf-8").read().splitlines()
i, para = 0, []


def flush():
    global para
    if para:
        text = " ".join(para).strip()
        if text:
            p = doc.add_paragraph()
            add_runs(p, text)
            if re.match(r"\*\*(Figure|Table)\s", text):
                for r in p.runs:
                    r.font.size = Pt(9)
        para = []


while i < len(lines):
    s = lines[i].strip()
    if not s:
        flush(); i += 1; continue
    if s.startswith("!["):
        flush()
        m = re.search(r"\(([^)]+)\)", s)
        if m:
            add_image(m.group(1))
        i += 1; continue
    if s.startswith("|"):
        flush()
        block = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            block.append(lines[i]); i += 1
        add_table(block); continue
    if s.startswith("### "):
        flush(); doc.add_heading(s[4:], level=3); i += 1; continue
    if s.startswith("## "):
        flush(); doc.add_heading(s[3:], level=1); i += 1; continue
    if s.startswith("# "):
        flush(); doc.add_heading(s[2:], level=0); i += 1; continue
    if s.startswith("---"):
        flush(); i += 1; continue
    if s.startswith("> "):
        flush()
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        add_runs(p, s[2:])
        for r in p.runs:
            r.italic = True; r.font.size = Pt(9)
        i += 1; continue
    if re.match(r"^(\-|\d+\.)\s", s):
        flush()
        p = doc.add_paragraph(style="List Bullet" if s[0] == "-" else "List Number")
        add_runs(p, re.sub(r"^(\-|\d+\.)\s", "", s))
        i += 1; continue
    para.append(s); i += 1

flush()
doc.save(OUT)
print("wrote", OUT, "(%d KB)" % (os.path.getsize(OUT) // 1024))
